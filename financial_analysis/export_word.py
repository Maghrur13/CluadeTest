"""
Export financial data and ratios to a formatted Word document.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def export_to_word(
    raw_data: dict,
    ratios: dict,
    output_path: str | Path,
    company_name: str = "Company",
) -> Path:
    """
    Create a formatted Word document with financial analysis.
    """
    output_path = Path(output_path)
    doc = Document()

    # Title
    title = doc.add_heading(f"Financial Statement Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(company_name).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Section 1: Extracted Data
    doc.add_heading("1. Extracted Financial Data", level=1)
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = "Table Grid"
    hdr = table1.rows[0].cells
    hdr[0].text = "Line Item"
    hdr[1].text = "Value"
    for key, val in raw_data.items():
        if key.startswith("_"):
            continue
        row = table1.add_row()
        row.cells[0].text = key.replace("_", " ").title()
        row.cells[1].text = str(val) if not isinstance(val, (int, float)) else f"{val:,.2f}"
    doc.add_paragraph()

    # Section 2: Ratio Analysis
    doc.add_heading("2. Financial Ratio Analysis", level=1)

    for category, items in ratios.items():
        doc.add_heading(category, level=2)
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = "Table Grid"
        h = table2.rows[0].cells
        h[0].text = "Ratio"
        h[1].text = "Value"
        h[2].text = "Interpretation"
        interpretations = {
            "Current Ratio": ">1.5 healthy, <1 may indicate liquidity risk",
            "Quick Ratio (Acid Test)": ">1 preferred; excludes inventory",
            "Cash Ratio": "Most conservative liquidity measure",
            "Debt-to-Equity": "Lower = less leverage; varies by industry",
            "Debt-to-Assets": "Lower = less debt-financed",
            "Interest Coverage": ">2 healthy; ability to pay interest",
            "Net Profit Margin (%)": "Higher = more profitable per dollar of sales",
            "Return on Assets (ROA) (%)": "Efficiency of asset use",
            "Return on Equity (ROE) (%)": "Return to shareholders",
            "Gross Margin (%)": "Profit after direct costs",
            "Operating Margin (%)": "Profit from core operations",
        }
        for ratio_name, value in items.items():
            row = table2.add_row()
            row.cells[0].text = ratio_name
            row.cells[1].text = f"{value:,.2f}" if isinstance(value, (int, float)) else str(value)
            row.cells[2].text = interpretations.get(ratio_name, "")
        doc.add_paragraph()

    doc.save(output_path)
    return output_path
